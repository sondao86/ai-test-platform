"""DOCX document parser using python-docx."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


def parse_docx(file_path: str | Path) -> str:
    """Extract text content from a DOCX file.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        Extracted text as a single string.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    doc = Document(str(file_path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Preserve heading structure
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            prefix = "#" * int(level) if level.isdigit() else "#"
            paragraphs.append(f"{prefix} {text}")
        else:
            paragraphs.append(text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n[Table]\n" + "\n".join(rows))

    full_text = "\n\n".join(paragraphs)
    logger.info("Parsed DOCX %s: %d paragraphs, %d characters", file_path.name, len(paragraphs), len(full_text))
    return full_text
